from __future__ import annotations

import json
from pathlib import Path
from functools import wraps
from urllib.parse import urlparse
import sqlite3
from datetime import datetime
import random
import io
import re
import textwrap
from flask import Flask, render_template, request, redirect, url_for, flash, session, g, jsonify, send_file
from werkzeug.security import generate_password_hash, check_password_hash

try:
    import markdown as md  # type: ignore
except Exception:  # pragma: no cover
    md = None  # type: ignore
try:
    import bleach  # type: ignore
except Exception:  # pragma: no cover
    bleach = None  # type: ignore
from markupsafe import escape
from concurrent.futures import ThreadPoolExecutor

from config import get_settings
from services.llm_client import LLMClient, build_prompt, parse_sections
from services.i18n import t, get_supported_languages, resolve_lang_from_request_args, language_name
from services import db as db_service

# --- Markdown rendering (safe) ---
ALLOWED_TAGS = [
    'p', 'a', 'strong', 'em', 'ul', 'ol', 'li', 'code', 'pre', 'blockquote', 'hr', 'br',
    'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'table', 'thead', 'tbody', 'tr', 'th', 'td'
]
ALLOWED_ATTRS = {
    'a': ['href', 'title', 'rel', 'target'],
    'code': ['class'],
    'th': ['colspan', 'rowspan'],
    'td': ['colspan', 'rowspan'],
}

def render_markdown_safe(text: str) -> str:
    if not text:
        return ''
    # Fallback if packages are unavailable at runtime
    if md is None or bleach is None:
        # Escape HTML and convert newlines to <br>
        return str(escape(text)).replace('\n', '<br>')
    # Convert Markdown to HTML
    html = md.markdown(text, extensions=['extra', 'sane_lists', 'nl2br'])
    # Sanitize HTML
    clean = bleach.clean(html, tags=ALLOWED_TAGS, attributes=ALLOWED_ATTRS, strip=True)
    # Auto-linkify and set safe link attributes
    def _target_blank(attrs, new=False):
        href_key = (None, 'href')
        if href_key in attrs:
            # open external links in new tab and ensure safety
            attrs[(None, 'target')] = '_blank'
            existing_rel = attrs.get((None, 'rel'), '') or ''
            rel_vals = set(existing_rel.split()) if existing_rel else set()
            rel_vals.update(['nofollow', 'noopener', 'noreferrer'])
            attrs[(None, 'rel')] = ' '.join(sorted(rel_vals))
        return attrs
    callbacks = list(getattr(bleach.linkifier, 'DEFAULT_CALLBACKS', []))
    callbacks.append(_target_blank)
    clean = bleach.linkify(clean, callbacks=callbacks)
    return clean

app = Flask(__name__)
settings = get_settings()
app.config['SECRET_KEY'] = settings.flask_secret_key

# Fields expected in the form/JSON
FORM_FIELDS = [
    'name', 'age', 'sex', 'chief_complaint', 'history', 'red_flags', 'vitals', 'goals',
    'biomechanical', 'respiratory', 'metabolic', 'neurological', 'behavioral'
]

BASE_DIR = Path(__file__).resolve().parent
DEMO_DIR = BASE_DIR / 'demo-data'


def get_db() -> sqlite3.Connection:
    return db_service.get_db()


def close_db(e=None):  # noqa: unused-argument
    return db_service.close_db(e)


@app.teardown_appcontext
def teardown_db(exception):  # noqa: unused-argument
    close_db()


# --- Background processing queue ---
executor = ThreadPoolExecutor(max_workers=settings.queue_max_concurrency)


def _new_db_conn() -> sqlite3.Connection:
    """Create a new standalone DB connection (not tied to Flask g)."""
    return db_service._new_db_conn()


def _resolve_model_for_provider(provider: str) -> str:
    if provider == 'OLLAMA':
        return settings.ollama_model or 'llama3.1'
    if provider == 'OPENAI':
        return settings.openai_model or 'gpt-4o-mini'
    if provider == 'GEMINI':
        return settings.gemini_model or 'gemini-1.5-flash'
    return ''


def create_job(payload: dict, lang: str, user: str | None) -> int:
    provider = settings.llm_provider
    model = _resolve_model_for_provider(provider)
    prompt = build_prompt(payload, lang=lang)
    now = datetime.utcnow().isoformat(timespec='seconds')
    db = get_db()
    cur = db.execute(
        """
        INSERT INTO jobs (created_at, user, lang, status, provider, model, payload_json, prompt_text)
        VALUES (?, ?, ?, 'queued', ?, ?, ?, ?)
        """,
        (
            now,
            user,
            lang,
            provider,
            model,
            json.dumps(payload, ensure_ascii=False),
            prompt,
        ),
    )
    db.commit()
    job_id = cur.lastrowid
    # Submit to executor
    executor.submit(process_job, job_id)
    return int(job_id)


def process_job(job_id: int) -> None:
    conn = _new_db_conn()
    try:
        row = conn.execute("SELECT * FROM jobs WHERE id = ?", (job_id,)).fetchone()
        if not row:
            return
        # Mark as running
        conn.execute(
            "UPDATE jobs SET status = 'running', started_at = ? WHERE id = ?",
            (datetime.utcnow().isoformat(timespec='seconds'), job_id),
        )
        conn.commit()
        prompt = row['prompt_text']
        lang = row['lang']
        provider = row['provider']
        model = row['model']
        payload_json = row['payload_json']
        user = row['user']

        client = LLMClient()
        resp = client.generate(prompt)
        if not resp.ok:
            conn.execute(
                "UPDATE jobs SET status='error', error_text=?, finished_at=? WHERE id=?",
                (str(resp.error), datetime.utcnow().isoformat(timespec='seconds'), job_id),
            )
            conn.commit()
            return

        # Persist to history
        cur = conn.execute(
            """
            INSERT INTO history (created_at, user, lang, provider, model, payload_json, prompt_text, response_text)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                datetime.utcnow().isoformat(timespec='seconds'),
                user,
                lang,
                provider,
                model,
                payload_json,
                prompt,
                resp.text,
            ),
        )
        history_id = cur.lastrowid
        conn.execute(
            "UPDATE jobs SET status='done', response_text=?, history_id=?, finished_at=? WHERE id=?",
            (resp.text, history_id, datetime.utcnow().isoformat(timespec='seconds'), job_id),
        )
        conn.commit()
    except Exception as e:
        try:
            conn.execute(
                "UPDATE jobs SET status='error', error_text=?, finished_at=? WHERE id=?",
                (str(e), datetime.utcnow().isoformat(timespec='seconds'), job_id),
            )
            conn.commit()
        except Exception:
            pass
    finally:
        conn.close()


def init_db():
    return db_service.init_db()


# Flask 3 removed before_first_request; initialize DB on the first incoming request instead.
_db_init_done = False

@app.before_request
def _ensure_db():
    global _db_init_done
    if not _db_init_done:
        init_db()
        _db_init_done = True


def _has_users() -> bool:
    try:
        db = get_db()
        row = db.execute("SELECT COUNT(*) AS c FROM users").fetchone()
        return (row[0] if isinstance(row, tuple) else row["c"]) > 0
    except Exception:
        return False


def _auth_enabled() -> bool:
    # Enable auth if at least one user exists in DB
    return _has_users()


def _is_logged_in() -> bool:
    return bool(session.get('auth', False))


def get_current_user():
    if not _is_logged_in():
        return None
    user_id = session.get('user_id')
    if not user_id:
        return None
    try:
        db = get_db()
        row = db.execute("SELECT id, username, is_admin, is_active, created_at FROM users WHERE id = ?", (user_id,)).fetchone()
        return row
    except Exception:
        return None


def admin_required(view_func):
    @wraps(view_func)
    def wrapper(*args, **kwargs):
        if not _auth_enabled():
            return redirect(url_for('index', lang=resolve_lang_from_request_args(request.args)))
        user = get_current_user()
        if user and user['is_admin']:
            return view_func(*args, **kwargs)
        flash('Admin access required.', 'warning')
        return redirect(url_for('index', lang=resolve_lang_from_request_args(request.args)))
    return wrapper


@app.context_processor
def inject_user():
    return {
        'current_user': get_current_user(),
        'is_logged_in': _is_logged_in(),
        'auth_enabled': _auth_enabled(),
    }


def login_required(view_func):
    @wraps(view_func)
    def wrapper(*args, **kwargs):
        if not _auth_enabled():
            return view_func(*args, **kwargs)
        if _is_logged_in():
            return view_func(*args, **kwargs)
        # Preserve lang and next
        lang = resolve_lang_from_request_args(request.args)
        next_url = request.full_path if request.query_string else request.path
        return redirect(url_for('login', next=next_url, lang=lang))
    return wrapper


def _is_safe_next(next_url: str | None) -> bool:
    if not next_url:
        return False
    # Disallow absolute URLs to prevent open redirects
    parsed = urlparse(next_url)
    return not parsed.netloc and (parsed.path or parsed.query)


@app.route('/login', methods=['GET', 'POST'])
def login():
    lang = resolve_lang_from_request_args(request.args if request.method == 'GET' else request.form)
    def tn(key):
        return t(key, lang)

    # Helper to generate a simple math captcha
    def _new_captcha() -> str:
        a = random.randint(1, 9)
        b = random.randint(1, 9)
        session['captcha_answer'] = str(a + b)
        return f"{a} + {b} = ?"

    if not _auth_enabled():
        # If no users yet, redirect to registration so user can create the first account
        return redirect(url_for('register', lang=lang))

    captcha_question = None

    if request.method == 'POST':
        username = (request.form.get('username') or '').strip()
        password = request.form.get('password') or ''
        captcha = (request.form.get('captcha') or '').strip()
        if not captcha:
            flash(tn('err.captcha_required'), 'warning')
        elif captcha != session.get('captcha_answer'):
            flash(tn('err.captcha_invalid'), 'danger')
        else:
            db = get_db()
            row = db.execute("SELECT id, username, password_hash, is_admin, is_active FROM users WHERE username = ?", (username,)).fetchone()
            if not row or not check_password_hash(row['password_hash'], password):
                flash(tn('err.login_invalid'), 'danger')
            elif not row['is_active']:
                flash(tn('err.login_blocked'), 'warning')
            else:
                session.clear()
                session['auth'] = True
                session['user'] = row['username']
                session['user_id'] = row['id']
                session['is_admin'] = bool(row['is_admin'])
                next_url = request.args.get('next') or request.form.get('next')
                if _is_safe_next(next_url):
                    return redirect(next_url)
                return redirect(url_for('index', lang=lang))
        # regenerate captcha after any POST error
        captcha_question = _new_captcha()
        # fall through to render template on error
    else:
        # GET: generate captcha to display
        captcha_question = _new_captcha()

    next_url = request.args.get('next') if _is_safe_next(request.args.get('next')) else url_for('index', lang=lang)
    return render_template('login.html', settings=settings, lang=lang, langs=get_supported_languages(), tn=tn, next_url=next_url, captcha_question=captcha_question)


@app.route('/register', methods=['GET', 'POST'])
def register():
    # Public self-registration: first user becomes admin; others are regular users
    lang = resolve_lang_from_request_args(request.args if request.method == 'GET' else request.form)
    def tn(key):
        return t(key, lang)

    # Helper to generate a simple math captcha
    def _new_captcha() -> str:
        a = random.randint(1, 9)
        b = random.randint(1, 9)
        session['captcha_answer'] = str(a + b)
        return f"{a} + {b} = ?"

    captcha_question = None

    if request.method == 'POST':
        username = (request.form.get('username') or '').strip()
        password = request.form.get('password') or ''
        password2 = request.form.get('password2') or ''
        captcha = (request.form.get('captcha') or '').strip()
        if not captcha:
            flash(tn('err.captcha_required'), 'warning')
        elif captcha != session.get('captcha_answer'):
            flash(tn('err.captcha_invalid'), 'danger')
        elif not username or not password:
            flash('Username and password are required.', 'warning')
        elif password != password2:
            flash('Passwords do not match.', 'warning')
        elif len(password) < 6:
            flash('Password must be at least 6 characters.', 'warning')
        else:
            db = get_db()
            existing = db.execute("SELECT id FROM users WHERE username = ?", (username,)).fetchone()
            if existing:
                flash('Username is already taken.', 'warning')
            else:
                is_first = not _has_users()
                is_admin = 1 if (is_first or session.get('is_admin')) else 0
                db.execute(
                    "INSERT INTO users (username, password_hash, is_admin, is_active, created_at) VALUES (?, ?, ?, 1, ?)",
                    (
                        username,
                        generate_password_hash(password),
                        is_admin,
                        datetime.utcnow().isoformat(timespec='seconds'),
                    ),
                )
                db.commit()
                # Auto-login new user
                row = db.execute("SELECT id, username, is_admin FROM users WHERE username = ?", (username,)).fetchone()
                session.clear()
                session['auth'] = True
                session['user'] = row['username']
                session['user_id'] = row['id']
                session['is_admin'] = bool(row['is_admin'])
                flash('Account created successfully.', 'success')
                return redirect(url_for('index', lang=lang))
        # regenerate captcha after any POST error
        captcha_question = _new_captcha()
    else:
        # GET: generate captcha to display
        captcha_question = _new_captcha()
    return render_template('register.html', settings=settings, lang=lang, langs=get_supported_languages(), tn=tn, captcha_question=captcha_question)


@app.route('/logout', methods=['POST', 'GET'])
def logout():
    lang = resolve_lang_from_request_args(request.args if request.method == 'GET' else request.form)
    session.clear()
    return redirect(url_for('login', lang=lang))


def _filter_payload(data: dict) -> dict:
    payload: dict[str, str] = {}
    for key in FORM_FIELDS:
        val = data.get(key, '') if isinstance(data, dict) else ''
        if isinstance(val, (int, float)):
            val = str(val)
        elif not isinstance(val, str):
            val = ''
        payload[key] = val
    return payload


def _list_demo_files() -> list[dict]:
    demos = []
    if DEMO_DIR.exists():
        for p in sorted(DEMO_DIR.glob('*.json')):
            demos.append({'name': p.stem, 'filename': p.name})
    return demos


@app.route('/', methods=['GET'])
@login_required
def index():
    lang = resolve_lang_from_request_args(request.args)
    demos = _list_demo_files()
    # translator lambda bound to lang
    def tn(key):
        return t(key, lang)
    return render_template('index.html', settings=settings, payload={}, demos=demos, lang=lang, langs=get_supported_languages(), tn=tn)


@app.route('/load-json', methods=['GET', 'POST'])
@login_required
def load_json():
    try:
        lang = resolve_lang_from_request_args(request.args if request.method == 'GET' else request.form)
        demos = _list_demo_files()
        def tn(key):
            return t(key, lang)
        if request.method == 'GET':
            demo = request.args.get('demo')
            if demo:
                candidate = (DEMO_DIR / f"{Path(demo).stem}.json").resolve()
                # Ensure the path is inside demo dir
                if not str(candidate).startswith(str(DEMO_DIR.resolve())) or not candidate.exists():
                    flash(tn('err.demo_not_found'), 'warning')
                    return redirect(url_for('index', lang=lang))
                with candidate.open('r', encoding='utf-8') as f:
                    data = json.load(f)
                payload = _filter_payload(data)
                return render_template('index.html', settings=settings, payload=payload, demos=demos, lang=lang, langs=get_supported_languages(), tn=tn)
            # If no demo specified, just redirect
            return redirect(url_for('index', lang=lang))

        # POST: file upload
        file = request.files.get('json_file')
        if not file or file.filename == '':
            flash(tn('err.json_choose'), 'warning')
            return redirect(url_for('index', lang=lang))
        try:
            raw = file.read().decode('utf-8')
            data = json.loads(raw)
        except Exception:
            flash(tn('err.json_invalid'), 'danger')
            return redirect(url_for('index', lang=lang))
        payload = _filter_payload(data)
        return render_template('index.html', settings=settings, payload=payload, demos=demos, lang=lang, langs=get_supported_languages(), tn=tn)
    except Exception as e:
        flash(f"{t('err.load_json', lang)} {e}", 'danger')
        return redirect(url_for('index', lang=resolve_lang_from_request_args(request.args)))


@app.route('/save', methods=['POST'])
@login_required
def save_draft():
    form = request.form
    lang = resolve_lang_from_request_args(form)
    payload = {
        'name': form.get('name', '').strip(),
        'age': form.get('age', '').strip(),
        'sex': form.get('sex', '').strip(),
        'chief_complaint': form.get('chief_complaint', '').strip(),
        'history': form.get('history', '').strip(),
        'red_flags': form.get('red_flags', '').strip(),
        'vitals': form.get('vitals', '').strip(),
        'goals': form.get('goals', '').strip(),
        'biomechanical': form.get('biomechanical', '').strip(),
        'respiratory': form.get('respiratory', '').strip(),
        'metabolic': form.get('metabolic', '').strip(),
        'neurological': form.get('neurological', '').strip(),
        'behavioral': form.get('behavioral', '').strip(),
    }
    db = get_db()
    now = datetime.utcnow().isoformat(timespec='seconds')
    draft_id = form.get('draft_id')
    user = session.get('user')

    def tn(key):
        return t(key, lang)

    saved_id = None
    if draft_id:
        # Ownership check when auth is enabled
        if _auth_enabled() and not session.get('is_admin'):
            row = db.execute("SELECT id, user FROM drafts WHERE id=?", (int(draft_id),)).fetchone()
            if not row or (row['user'] and row['user'] != user):
                flash(tn('err.draft_not_found'), 'warning')
                demos = _list_demo_files()
                return render_template('index.html', settings=settings, payload=payload, demos=demos, lang=lang, langs=get_supported_languages(), tn=tn)
        try:
            db.execute(
                "UPDATE drafts SET updated_at=?, user=?, lang=?, payload_json=? WHERE id=?",
                (now, user, lang, json.dumps(payload, ensure_ascii=False), int(draft_id))
            )
            db.commit()
            saved_id = int(draft_id)
        except Exception:
            saved_id = None
    else:
        cur = db.execute(
            "INSERT INTO drafts (created_at, updated_at, user, lang, payload_json) VALUES (?, ?, ?, ?, ?)",
            (now, now, user, lang, json.dumps(payload, ensure_ascii=False))
        )
        db.commit()
        saved_id = cur.lastrowid

    if saved_id:
        flash(tn('flash.draft_saved') + f" #{saved_id}", 'success')
    else:
        flash(tn('flash.draft_saved'), 'success')
    demos = _list_demo_files()
    return render_template('index.html', settings=settings, payload=payload, demos=demos, lang=lang, langs=get_supported_languages(), tn=tn, draft_id=saved_id)


@app.route('/draft/<int:draft_id>', methods=['GET'])
@login_required
def open_draft(draft_id: int):
    lang = resolve_lang_from_request_args(request.args)
    db = get_db()
    row = db.execute("SELECT * FROM drafts WHERE id=?", (draft_id,)).fetchone()
    def tn(key):
        return t(key, lang)
    if not row:
        flash(tn('err.draft_not_found'), 'warning')
        return redirect(url_for('index', lang=lang))
    # Ownership enforcement when auth is enabled
    if _auth_enabled() and not session.get('is_admin'):
        owner = row['user']
        if owner and owner != session.get('user'):
            flash(tn('err.draft_not_found'), 'warning')
            return redirect(url_for('index', lang=lang))
    try:
        payload = json.loads(row['payload_json']) if row['payload_json'] else {}
    except Exception:
        payload = {}
    demos = _list_demo_files()
    return render_template('index.html', settings=settings, payload=payload, demos=demos, lang=lang, langs=get_supported_languages(), tn=tn, draft_id=row['id'])


@app.route('/draft/<int:draft_id>/delete', methods=['POST'])
@login_required
def delete_draft(draft_id: int):
    lang = resolve_lang_from_request_args(request.args if request.method == 'GET' else request.form)
    db = get_db()
    row = db.execute("SELECT id, user FROM drafts WHERE id=?", (draft_id,)).fetchone()
    def tn(key):
        return t(key, lang)
    if not row:
        flash(tn('err.draft_not_found'), 'warning')
        return redirect(url_for('history', lang=lang))
    if _auth_enabled() and not session.get('is_admin'):
        owner = row['user']
        if owner and owner != session.get('user'):
            flash(tn('err.draft_not_found'), 'warning')
            return redirect(url_for('history', lang=lang))
    db.execute("DELETE FROM drafts WHERE id=?", (draft_id,))
    db.commit()
    flash(tn('flash.draft_deleted'), 'success')
    return redirect(url_for('history', lang=lang))


@app.route('/analyze', methods=['POST'])
@login_required
def analyze():
    form = request.form
    lang = resolve_lang_from_request_args(form)
    payload = {
        'name': form.get('name', '').strip(),
        'age': form.get('age', '').strip(),
        'sex': form.get('sex', '').strip(),
        'chief_complaint': form.get('chief_complaint', '').strip(),
        'history': form.get('history', '').strip(),
        'red_flags': form.get('red_flags', '').strip(),
        'vitals': form.get('vitals', '').strip(),
        'goals': form.get('goals', '').strip(),
        'biomechanical': form.get('biomechanical', '').strip(),
        'respiratory': form.get('respiratory', '').strip(),
        'metabolic': form.get('metabolic', '').strip(),
        'neurological': form.get('neurological', '').strip(),
        'behavioral': form.get('behavioral', '').strip(),
    }

    def tn(key):
        return t(key, lang)

    if not payload['chief_complaint']:
        flash(tn('err.chief_required'), 'warning')
        return redirect(url_for('index', lang=lang))

    # Enqueue background job and redirect to status page
    job_id = create_job(payload, lang, session.get('user'))
    return redirect(url_for('job_status', job_id=job_id, lang=lang))


@app.route('/job/<int:job_id>', methods=['GET'])
@login_required
def job_status(job_id: int):
    lang = resolve_lang_from_request_args(request.args)
    def tn(key):
        return t(key, lang)
    db = get_db()
    row = db.execute("SELECT * FROM jobs WHERE id = ?", (job_id,)).fetchone()
    if not row:
        flash('Job not found.', 'warning')
        return redirect(url_for('index', lang=lang))
    return render_template('job_status.html', job=row, settings=settings, lang=lang, langs=get_supported_languages(), tn=tn)


@app.route('/api/job/<int:job_id>/status', methods=['GET'])
@login_required
def job_status_api(job_id: int):
    db = get_db()
    row = db.execute("SELECT status, error_text, history_id, lang FROM jobs WHERE id = ?", (job_id,)).fetchone()
    if not row:
        return jsonify({"ok": False, "error": "Job not found"}), 404
    return jsonify({
        "ok": True,
        "status": row['status'],
        "error_text": row['error_text'],
        "history_id": row['history_id'],
        "lang": row['lang'],
    })


@app.route('/result/<int:history_id>', methods=['GET'])
@login_required
def result_page(history_id: int):
    # Keep UI language from query param; content is from stored response
    lang = resolve_lang_from_request_args(request.args)
    def tn(key):
        return t(key, lang)
    db = get_db()
    row = db.execute("SELECT * FROM history WHERE id = ?", (history_id,)).fetchone()
    if not row:
        flash('History item not found.', 'warning')
        return redirect(url_for('history', lang=lang))
    try:
        payload = json.loads(row['payload_json']) if row['payload_json'] else {}
    except Exception:
        payload = {}
    resp_text = row['response_text'] or ''
    sections = parse_sections(resp_text)
    sections_html = {
        'summary': render_markdown_safe(sections.get('summary', '')),
        'diagnosis': render_markdown_safe(sections.get('diagnosis', '')),
        'plan': render_markdown_safe(sections.get('plan', '')),
    }
    return render_template('result.html', payload=payload, sections=sections, sections_html=sections_html, raw=resp_text, history_id=history_id, settings=settings, lang=lang, langs=get_supported_languages(), tn=tn)


@app.route('/result/<int:history_id>/download.pdf', methods=['GET'])
@login_required
def result_download_pdf(history_id: int):
    lang = resolve_lang_from_request_args(request.args)
    db = get_db()
    row = db.execute("SELECT * FROM history WHERE id = ?", (history_id,)).fetchone()
    if not row:
        flash('History item not found.', 'warning')
        return redirect(url_for('history', lang=lang))
    try:
        payload = json.loads(row['payload_json']) if row['payload_json'] else {}
    except Exception:
        payload = {}
    # Try to import reportlab
    try:
        from reportlab.pdfgen import canvas  # type: ignore
        from reportlab.lib.pagesizes import A4  # type: ignore
        from reportlab.lib.units import cm  # type: ignore
    except Exception:
        flash('PDF export is not available on this server (missing reportlab).', 'warning')
        return redirect(url_for('result_page', history_id=history_id, lang=lang))

    sections = parse_sections(row['response_text'] or '')

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4

    def slug(s: str) -> str:
        s = (s or '').lower()
        s = re.sub(r'[^a-z0-9]+', '-', s)
        s = s.strip('-')
        return s or 'patient'

    # Compose content
    margin = 2 * cm
    text = c.beginText(margin, height - margin)
    text.setFont("Helvetica-Bold", 16)
    title = "OsteoDiag Results"
    name = payload.get('name') or 'Patient'
    text.textLine(f"{title} — {name}")

    text.setFont("Helvetica", 10)
    meta = f"Provider: {row['provider']}  Model: {row['model']}  Lang: {row['lang']}  Created: {row['created_at']}"
    text.textLine(meta)
    text.textLine(" ")

    def write_section(heading: str, body: str):
        text.setFont("Helvetica-Bold", 12)
        text.textLine(heading)
        text.setFont("Helvetica", 11)
        if not body:
            text.textLine("-")
            text.textLine(" ")
            return
        wrapped = textwrap.fill(body, width=95)
        for line in wrapped.splitlines():
            text.textLine(line)
        text.textLine(" ")

    write_section("Summary", sections.get('summary', ''))
    write_section("Diagnosis", sections.get('diagnosis', ''))
    write_section("Treatment Plan", sections.get('plan', ''))

    c.drawText(text)
    c.showPage()
    c.save()
    buf.seek(0)

    filename = f"{slug(payload.get('name','')) or 'patient'}-result-{history_id}.pdf"
    return send_file(buf, mimetype='application/pdf', as_attachment=True, download_name=filename)


@app.route('/result/<int:history_id>/download.docx', methods=['GET'])
@login_required
def result_download_docx(history_id: int):
    lang = resolve_lang_from_request_args(request.args)
    db = get_db()
    row = db.execute("SELECT * FROM history WHERE id = ?", (history_id,)).fetchone()
    if not row:
        flash('History item not found.', 'warning')
        return redirect(url_for('history', lang=lang))
    try:
        payload = json.loads(row['payload_json']) if row['payload_json'] else {}
    except Exception:
        payload = {}

    try:
        from docx import Document  # type: ignore
    except Exception:
        flash('DOCX export is not available on this server (missing python-docx).', 'warning')
        return redirect(url_for('result_page', history_id=history_id, lang=lang))

    sections = parse_sections(row['response_text'] or '')

    doc = Document()
    title = doc.add_heading('OsteoDiag Results', level=0)
    name = payload.get('name') or 'Patient'
    doc.add_paragraph(f"Patient: {name}")
    doc.add_paragraph(f"Provider: {row['provider']}  Model: {row['model']}  Lang: {row['lang']}  Created: {row['created_at']}")

    def add_section(h: str, body: str):
        doc.add_heading(h, level=1)
        if body:
            for para in body.split('\n\n'):
                doc.add_paragraph(para)
        else:
            doc.add_paragraph('-')

    add_section('Summary', sections.get('summary', ''))
    add_section('Diagnosis', sections.get('diagnosis', ''))
    add_section('Treatment Plan', sections.get('plan', ''))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    def slug(s: str) -> str:
        s = (s or '').lower()
        s = re.sub(r'[^a-z0-9]+', '-', s)
        s = s.strip('-')
        return s or 'patient'

    filename = f"{slug(payload.get('name','')) or 'patient'}-result-{history_id}.docx"
    return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, download_name=filename)


@app.route('/history/<int:history_id>/download.json', methods=['GET'])
@login_required
def history_download_json(history_id: int):
    lang = resolve_lang_from_request_args(request.args)
    db = get_db()
    row = db.execute("SELECT * FROM history WHERE id = ?", (history_id,)).fetchone()
    if not row:
        flash('History item not found.', 'warning')
        return redirect(url_for('history', lang=lang))
    # Ownership enforcement when auth is enabled and user is not admin
    if _auth_enabled() and not session.get('is_admin'):
        if (row['user'] or None) != session.get('user'):
            flash('History item not found.', 'warning')
            return redirect(url_for('history', lang=lang))
    try:
        payload = json.loads(row['payload_json']) if row['payload_json'] else {}
    except Exception:
        payload = {}

    def slug(s: str) -> str:
        s = (s or '').lower()
        s = re.sub(r'[^a-z0-9]+', '-', s)
        s = s.strip('-')
        return s or 'patient'

    out = io.BytesIO()
    out.write(json.dumps(payload, ensure_ascii=False, indent=2).encode('utf-8'))
    out.seek(0)
    filename = f"{slug(payload.get('name','')) or 'patient'}-payload-{history_id}.json"
    return send_file(out, mimetype='application/json', as_attachment=True, download_name=filename)


@app.route('/history', methods=['GET'])
@login_required
def history():
    lang = resolve_lang_from_request_args(request.args)
    def tn(key):
        return t(key, lang)
    db = get_db()
    # Pagination
    try:
        page = int(request.args.get('page', 1) or 1)
    except Exception:
        page = 1
    page = max(1, page)
    page_size = 20
    offset = (page - 1) * page_size

    # Filter completed examinations by user when auth is enabled and user is not admin
    is_admin = bool(session.get('is_admin'))
    where_sql = ""
    where_params = []
    if _auth_enabled() and not is_admin:
        where_sql = " WHERE user = ?"
        where_params = [session.get('user')]

    total_row = db.execute(f"SELECT COUNT(*) AS c FROM history{where_sql}", tuple(where_params)).fetchone()
    total = (total_row[0] if isinstance(total_row, tuple) else total_row['c'])
    pages = (total + page_size - 1) // page_size if total else 1

    rows = db.execute(
        f"SELECT id, created_at, user, lang, provider, model FROM history{where_sql} ORDER BY id DESC LIMIT ? OFFSET ?",
        tuple(where_params + [page_size, offset])
    ).fetchall()

    # Running processes (left unchanged per requirements)
    running = db.execute(
        "SELECT id, created_at, user, lang, provider, model, status FROM jobs WHERE status IN ('queued','running') ORDER BY id DESC"
    ).fetchall()

    # Drafts: only own drafts when auth enabled and not admin; otherwise all
    if _auth_enabled() and not is_admin:
        drafts = db.execute(
            "SELECT id, created_at, updated_at, user, lang FROM drafts WHERE user = ? ORDER BY updated_at DESC LIMIT 50",
            (session.get('user'),)
        ).fetchall()
    else:
        drafts = db.execute(
            "SELECT id, created_at, updated_at, user, lang FROM drafts ORDER BY updated_at DESC LIMIT 50"
        ).fetchall()

    return render_template('history.html', rows=rows, running=running, drafts=drafts, page=page, pages=pages, total=total, page_size=page_size, settings=settings, lang=lang, langs=get_supported_languages(), tn=tn)


@app.route('/history/<int:item_id>', methods=['GET'])
@login_required
def history_detail(item_id: int):
    lang = resolve_lang_from_request_args(request.args)
    def tn(key):
        return t(key, lang)
    db = get_db()
    row = db.execute(
        "SELECT * FROM history WHERE id = ?",
        (item_id,)
    ).fetchone()
    if not row:
        flash('History item not found.', 'warning')
        return redirect(url_for('history', lang=lang))
    # Ownership enforcement when auth is enabled and not admin
    if _auth_enabled() and not session.get('is_admin'):
        if (row['user'] or None) != session.get('user'):
            flash('History item not found.', 'warning')
            return redirect(url_for('history', lang=lang))
    return render_template('history_detail.html', item=row, settings=settings, lang=lang, langs=get_supported_languages(), tn=tn)


# --- Admin: users management ---
@app.route('/admin/users', methods=['GET'])
@admin_required
def admin_users():
    lang = resolve_lang_from_request_args(request.args)
    db = get_db()
    users = db.execute("SELECT id, username, is_admin, is_active, created_at FROM users ORDER BY id ASC").fetchall()
    def tn(key):
        return t(key, lang)
    return render_template('admin_users.html', users=users, settings=settings, lang=lang, langs=get_supported_languages(), tn=tn)


@app.route('/admin/users/<int:user_id>/block', methods=['POST'])
@admin_required
def admin_block_user(user_id: int):
    lang = resolve_lang_from_request_args(request.args if request.method == 'GET' else request.form)
    me = get_current_user()
    if me and me['id'] == user_id:
        flash("You can't block yourself.", 'warning')
        return redirect(url_for('admin_users', lang=lang))
    db = get_db()
    db.execute("UPDATE users SET is_active = 0 WHERE id = ?", (user_id,))
    db.commit()
    flash('User blocked.', 'success')
    return redirect(url_for('admin_users', lang=lang))


@app.route('/admin/users/<int:user_id>/unblock', methods=['POST'])
@admin_required
def admin_unblock_user(user_id: int):
    lang = resolve_lang_from_request_args(request.args if request.method == 'GET' else request.form)
    db = get_db()
    db.execute("UPDATE users SET is_active = 1 WHERE id = ?", (user_id,))
    db.commit()
    flash('User unblocked.', 'success')
    return redirect(url_for('admin_users', lang=lang))


@app.route('/admin/users/<int:user_id>/delete', methods=['POST'])
@admin_required
def admin_delete_user(user_id: int):
    lang = resolve_lang_from_request_args(request.args if request.method == 'GET' else request.form)
    me = get_current_user()
    if me and me['id'] == user_id:
        flash("You can't delete yourself.", 'warning')
        return redirect(url_for('admin_users', lang=lang))
    db = get_db()
    db.execute("DELETE FROM users WHERE id = ?", (user_id,))
    db.commit()
    flash('User deleted.', 'success')
    return redirect(url_for('admin_users', lang=lang))


if __name__ == '__main__':
    host = settings.host
    port = settings.port
    app.run(host=host, port=port, debug=settings.flask_debug)
